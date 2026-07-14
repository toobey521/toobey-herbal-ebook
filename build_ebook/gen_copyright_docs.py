#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 Toobey本草食蔬 软件著作权登记三件套
输出到桌面: 
  01-软件著作权登记申请表.docx
  02-软件说明书.docx  
  03-源代码文档.docx
"""
import os, sys, json
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DESKTOP = r'C:\Users\Administrator\Desktop'
BUILD_DIR = r'C:\Users\Administrator\Desktop\build_ebook'

# ── 中文字体设置 ──
def set_run_font(run, font_name='SimSun', size=11, bold=False, color=None, cn_font='SimSun'):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)
    if color:
        run.font.color.rgb = color

def add_heading_cn(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, 'SimHei', size=16 if level==1 else 14, bold=True, cn_font='SimHei')
    return h

def add_para(doc, text, size=11, bold=False, align=None, spacing_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 'SimSun', size=size, bold=bold, cn_font='SimSun')
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    return p

def add_table_row(table, cells_data, bold=False, size=10, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        set_run_font(run, 'SimSun', size=size, bold=bold or header, cn_font='SimSun')
    return row


# ═══════════════════════════════════════════════════
# 文档1：软件著作权登记申请表
# ═══════════════════════════════════════════════════
def gen_application_form():
    doc = Document()
    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    
    # ── 标题 ──
    add_heading_cn(doc, '计算机软件著作权登记申请表', level=1)
    add_para(doc, '', size=6)
    
    # ── 基本信息表 ──
    info_data = [
        ('软件名称', 'Toobey本草食蔬'),
        ('软件简称', '本草食蔬'),
        ('版本号', 'V1.0'),
        ('软件分类', '应用软件-信息查询软件'),
        ('开发完成日期', '2026年7月14日'),
        ('首次发表日期', '2026年7月14日'),
        ('权利取得方式', '原始取得'),
        ('权利范围', '全部权利'),
        ('著作权人', 'Toobey（个人）'),
        ('开发者', 'Toobey'),
        ('开发方式', '独立开发'),
        ('编程语言', 'Python 3.11'),
        ('源程序量', '约430行（核心逻辑）+ 319条食材数据 + 322张图片'),
        ('软件运行环境', 'Windows 10/11（64位）'),
        ('硬件要求', 'CPU 1.0GHz以上，内存2GB以上，硬盘100MB以上空间'),
        ('软件用途', '查阅各类蔬菜的性味归经、寒热属性、药用价值及营养成分等信息'),
    ]
    
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for label, value in info_data:
        add_table_row(table, [label, value], bold=False, size=10)
    
    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(12)
    
    add_para(doc, '', size=6)
    
    # ── 功能特点 ──
    add_heading_cn(doc, '一、软件功能与特点', level=2)
    
    features = [
        '收录319种常见蔬菜的完整信息，涵盖叶菜类、根茎类、瓜类、菌菇类、豆类、海藻类、水果类等',
        '每种蔬菜详列性味归经（寒/凉/平/温/热、甘/苦/辛/酸/咸）、功效主治、禁忌、营养成分和详细介绍',
        '按寒热属性分为五大类（寒性、凉性、平性、温性、热性），左侧导航栏快速切换',
        '支持关键词搜索（按菜名、功效、主治检索）',
        '每种蔬菜配有高清实物照片（共322张图片，来源于Bing图片搜索）',
        'GUI界面采用tkinter构建，左侧分类导航+右侧详细内容展示',
        '图片支持等比例缩放显示（最大宽280px高200px）',
        '右侧详情区域支持鼠标滚轮滚动',
        '单文件exe程序，双击即用，无需安装Python或任何依赖',
    ]
    for f in features:
        add_para(doc, f'• {f}', size=10, spacing_after=3)
    
    # ── 技术架构 ──
    add_heading_cn(doc, '二、技术架构', level=2)
    add_para(doc, '本软件采用Python 3.11 + tkinter桌面GUI框架开发。数据以JSON格式存储，程序启动时加载。图片文件打包在exe中，由PIL/Pillow库读取和缩放显示。单文件exe使用PyInstaller打包，内置所有资源和依赖。', size=10, spacing_after=6)
    
    # ── 数据来源 ──
    add_heading_cn(doc, '三、数据来源', level=2)
    add_para(doc, '蔬菜的性味归经、功效主治、禁忌等数据来源于中国传统中医药理论和《本草纲目》《食疗本草》等中医典籍整理。营养成分数据参考中国食物成分表。所有数据经过结构化整理和交叉校对。', size=10, spacing_after=6)
    
    # ── 声明 ──
    add_heading_cn(doc, '四、申请人声明', level=2)
    add_para(doc, '本人保证所提交的申请材料真实、完整、合法，并愿意承担因材料不实所产生的一切法律后果。', size=10)
    add_para(doc, '', size=20)
    add_para(doc, '申请人（签字/盖章）：______________', size=11, spacing_after=3)
    add_para(doc, '申请日期：______年______月______日', size=11, spacing_after=3)
    
    path = os.path.join(DESKTOP, '01-软件著作权登记申请表.docx')
    doc.save(path)
    print(f'✅ {path}')
    return path


# ═══════════════════════════════════════════════════
# 文档2：软件说明书
# ═══════════════════════════════════════════════════
def gen_manual():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    
    # ── 封面 ──
    add_para(doc, '', size=60)
    add_para(doc, 'Toobey本草食蔬', size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12)
    add_para(doc, '软件说明书', size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=40)
    add_para(doc, '', size=30)
    
    info_lines = [
        ('软件名称：', 'Toobey本草食蔬'),
        ('版本号：', 'V1.0'),
        ('开发者：', 'Toobey'),
        ('开发日期：', '2026年7月14日'),
        ('运行环境：', 'Windows 10/11'),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in info_lines:
        r = table.add_row()
        r.cells[0].text = ''
        r.cells[1].text = ''
        run0 = r.cells[0].paragraphs[0].add_run(label)
        set_run_font(run0, 'SimHei', 12, bold=True, cn_font='SimHei')
        run1 = r.cells[1].paragraphs[0].add_run(value)
        set_run_font(run1, 'SimSun', 12, cn_font='SimSun')
    
    doc.add_page_break()
    
    # ── 目录 ──
    add_heading_cn(doc, '目  录', level=1)
    toc_items = [
        '一、软件概述',
        '二、安装与运行',
        '三、功能说明',
        '  3.1 主界面布局',
        '  3.2 分类导航',
        '  3.3 蔬菜详情查看',
        '  3.4 搜索功能',
        '四、数据结构',
        '  4.1 蔬菜信息字段',
        '  4.2 寒热属性分类',
        '五、操作示例',
        '  5.1 按分类浏览',
        '  5.2 搜索特定蔬菜',
        '  5.3 查看详情',
        '六、常见问题',
    ]
    for item in toc_items:
        add_para(doc, item, size=11, spacing_after=3)
    
    doc.add_page_break()
    
    # ── 正文 ──
    # 一、软件概述
    add_heading_cn(doc, '一、软件概述', level=1)
    add_para(doc, 'Toobey本草食蔬是一款面向大众的蔬菜养生知识查询软件。软件收录了319种常见蔬菜的性味归经、寒热属性、功效主治、禁忌人群、营养成分和详细介绍等信息，是日常饮食养生和中医药知识查询的实用工具。', size=11, spacing_after=6)
    add_para(doc, '软件以简洁直观的GUI界面呈现，左侧为分类导航栏（按寒热属性分为五类），右侧为详细内容展示区域。支持关键词搜索和图片浏览，方便用户快速查找和了解各类蔬菜的药用价值和营养价值。', size=11, spacing_after=6)
    
    # 二、安装与运行
    add_heading_cn(doc, '二、安装与运行', level=1)
    add_para(doc, '本软件为绿色免安装软件，只需将"Toobey本草食蔬.exe"文件保存到电脑任意位置，双击即可运行。', size=11, spacing_after=3)
    add_para(doc, '系统要求：', size=11, bold=True, spacing_after=3)
    add_para(doc, '• 操作系统：Windows 10/11（64位）', size=10, spacing_after=2)
    add_para(doc, '• 处理器：1.0 GHz或以上', size=10, spacing_after=2)
    add_para(doc, '• 内存：2 GB或以上', size=10, spacing_after=2)
    add_para(doc, '• 硬盘空间：至少100 MB可用空间', size=10, spacing_after=2)
    add_para(doc, '• 无需安装Python、Java等运行环境', size=10, spacing_after=6)
    
    # 三、功能说明
    add_heading_cn(doc, '三、功能说明', level=1)
    
    add_heading_cn(doc, '3.1 主界面布局', level=2)
    add_para(doc, '主界面分为左右两栏布局：', size=11, spacing_after=3)
    add_para(doc, '• 左侧（约160px宽）：深色背景的分类导航栏，顶部显示"本草食蔬"标题，下方列出五个寒热属性分类及其包含的蔬菜数量。底部显示蔬菜总数。', size=10, spacing_after=2)
    add_para(doc, '• 右侧：白色背景的内容展示区，顶部有一个搜索框和搜索按钮，下方为蔬菜列表和详情展示区域。', size=10, spacing_after=2)
    
    add_heading_cn(doc, '3.2 分类导航', level=2)
    add_para(doc, '左侧导航栏将蔬菜按中医寒热属性分为五类：', size=11, spacing_after=3)
    add_para(doc, '❄️ 寒性（53种）：苦瓜、茭白、竹笋、海带、紫菜、鱼腥草等', size=10, spacing_after=2)
    add_para(doc, '🌊 凉性（89种）：白萝卜、丝瓜、冬瓜、芹菜、番茄、秋葵等', size=10, spacing_after=2)
    add_para(doc, '🌿 平性（96种）：白菜、山药、土豆、木耳、银耳、香菇等', size=10, spacing_after=2)
    add_para(doc, '🔥 温性（73种）：韭菜、大葱、大蒜、生姜、南瓜、洋葱等', size=10, spacing_after=2)
    add_para(doc, '⚡ 热性（6种）：辣椒、胡椒、肉桂、花椒等', size=10, spacing_after=2)
    add_para(doc, '点击任一分类，右侧列表即显示该分类下的所有蔬菜。', size=10, spacing_after=6)
    
    add_heading_cn(doc, '3.3 蔬菜详情查看', level=2)
    add_para(doc, '在右侧列表中点击任意蔬菜名称，详情区域即显示该蔬菜的完整信息：', size=11, spacing_after=3)
    add_para(doc, '• 蔬菜图片（顶部，等比例缩放显示，最大宽280px）', size=10, spacing_after=2)
    add_para(doc, '• 名称与性味标签（名称+属性标签+味型）', size=10, spacing_after=2)
    add_para(doc, '• 基本信息卡片：归经、功效、主治、禁忌、建议食法', size=10, spacing_after=2)
    add_para(doc, '• 营养成分（含具体数值和主要功效成分）', size=10, spacing_after=2)
    add_para(doc, '• 详细介绍（含食用方法、注意事项、文化典故等）', size=10, spacing_after=6)
    
    add_heading_cn(doc, '3.4 搜索功能', level=2)
    add_para(doc, '在右侧顶部的搜索框中输入关键词（蔬菜名称、功效如"清热"、主治如"咳嗽"等），点击"搜索"按钮，系统将在全部319种蔬菜中检索匹配项并显示结果。', size=11, spacing_after=6)
    
    # 四、数据结构
    add_heading_cn(doc, '四、数据结构', level=1)
    
    add_heading_cn(doc, '4.1 蔬菜信息字段', level=2)
    add_para(doc, '每条蔬菜记录包含以下10个字段：', size=11, spacing_after=3)
    fields_table = doc.add_table(rows=0, cols=3)
    fields_table.style = 'Table Grid'
    add_table_row(fields_table, ['字段名', '含义', '示例'], bold=True, size=10, header=True)
    field_rows = [
        ('name', '蔬菜名称', '苦瓜'),
        ('nature', '寒热属性', '寒/凉/平/温/热'),
        ('taste', '味型', '甘/苦/辛/酸/咸'),
        ('meridian', '归经', '归心、脾、肺经'),
        ('effect', '功效', '清热祛暑，明目解毒'),
        ('indications', '主治', '暑热烦渴、目赤肿痛'),
        ('contra', '禁忌', '脾胃虚寒者不宜多食'),
        ('usage', '建议食法', '凉拌、清炒、煲汤'),
        ('nutrition', '营养成分', '维生素C含量极高...'),
        ('desc', '详细介绍', '苦瓜虽苦，但苦味纯正...'),
        ('image', '图片文件名', '苦瓜.jpg'),
    ]
    for row_data in field_rows:
        add_table_row(fields_table, row_data, size=10)
    add_para(doc, '', size=6)
    
    add_heading_cn(doc, '4.2 寒热属性分类', level=2)
    add_para(doc, '软件中的所有蔬菜根据中医"四气"（寒热温凉）理论进行分类。每类蔬菜对应的颜色标识为：', size=11, spacing_after=3)
    add_para(doc, '• 寒性 → 蓝色（#3b82f6）', size=10, spacing_after=2)
    add_para(doc, '• 凉性 → 浅蓝色（#60a5fa）', size=10, spacing_after=2)
    add_para(doc, '• 平性 → 绿色（#22c55e）', size=10, spacing_after=2)
    add_para(doc, '• 温性 → 橙色（#f59e0b）', size=10, spacing_after=2)
    add_para(doc, '• 热性 → 红色（#ef4444）', size=10, spacing_after=6)
    
    # 五、操作示例
    add_heading_cn(doc, '五、操作示例', level=1)
    
    for ex in [
        ('5.1 按分类浏览', '双击桌面"Toobey本草食蔬.exe" → 软件自动打开 → 点击左侧"凉性"分类 → 右侧列表显示89种凉性蔬菜 → 点击"冬瓜" → 详情区显示冬瓜的图片、性味、功效、营养成分和详细介绍。'),
        ('5.2 搜索特定蔬菜', '在搜索框中输入"咳嗽" → 点击搜索 → 系统列出所有可用于治疗咳嗽的蔬菜（如白萝卜、百合、枇杷等） → 点击任一结果查看详情。'),
        ('5.3 查看图片', '点击任意蔬菜 → 详情区域顶部即显示该蔬菜的高清实物照片 → 图片自动缩放至合适大小。'),
    ]:
        add_heading_cn(doc, ex[0], level=2)
        add_para(doc, ex[1], size=11, spacing_after=6)
    
    # 六、常见问题
    add_heading_cn(doc, '六、常见问题', level=1)
    qas = [
        ('Q: 软件是否需要联网？', 'A: 不需要。所有数据和图片均已打包在exe文件中，离线即可使用。'),
        ('Q: 软件是否收费？', 'A: 本软件为免费软件，可以自由使用和传播。'),
        ('Q: 在Mac电脑上能否使用？', 'A: 当前版本仅支持Windows 10/11系统。'),
        ('Q: 数据是否有医学依据？', 'A: 蔬菜的性味归经和功效数据来源于传统中医理论和典籍整理，仅供参考，不能替代专业医疗建议。'),
    ]
    for q, a in qas:
        add_para(doc, q, size=11, bold=True, spacing_after=2)
        add_para(doc, a, size=11, spacing_after=6)
    
    path = os.path.join(DESKTOP, '02-软件说明书.docx')
    doc.save(path)
    print(f'✅ {path}')
    return path


# ═══════════════════════════════════════════════════
# 文档3：源代码文档
# ═══════════════════════════════════════════════════
def gen_source_code():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    
    add_heading_cn(doc, 'Toobey本草食蔬 源代码', level=1)
    add_para(doc, '软件版本：V1.0', size=11, spacing_after=3)
    add_para(doc, '编程语言：Python 3.11', size=11, spacing_after=3)
    add_para(doc, '核心代码文件：veggie_ebook.py', size=11, spacing_after=20)
    
    # ── 主程序源代码 ──
    src_path = os.path.join(BUILD_DIR, 'veggie_ebook.py')
    with open(src_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    total_lines = len(lines)
    add_para(doc, f'（全文共 {total_lines} 行）', size=10, spacing_after=12)
    add_heading_cn(doc, 'veggie_ebook.py 完整源代码', level=2)
    
    # 源代码用等宽字体展示
    code_block = ''
    for i, line in enumerate(lines, 1):
        # 替换特殊字符，避免docx问题
        text = line.rstrip('\n').replace('\t', '    ')
        code_block += f'{i:>4}| {text}\n'
    
    p = doc.add_paragraph()
    run = p.add_run(code_block)
    set_run_font(run, 'SimSun', size=7, cn_font='SimSun')
    p.paragraph_format.line_spacing = Pt(10)
    
    path = os.path.join(DESKTOP, '03-源代码文档.docx')
    doc.save(path)
    print(f'✅ {path}')
    return path


# ═══════════════════════════════════════════════════
# 主函数
# ═══════════════════════════════════════════════════
if __name__ == '__main__':
    print('开始生成软件著作权登记三件套...\n')
    print('=' * 50)
    
    p1 = gen_application_form()
    print()
    p2 = gen_manual()
    print()
    p3 = gen_source_code()
    
    print('\n' + '=' * 50)
    print('全部生成完成！✅')
    print(f'  📄 {p1}')
    print(f'  📄 {p2}')
    print(f'  📄 {p3}')
